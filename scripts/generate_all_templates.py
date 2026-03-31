#!/usr/bin/env python3
"""
生成 Office Pro Skill 的所有实际模板文件 (.docx 和 .xlsx)
"""

import os
import sys
import json
from pathlib import Path

SKILL_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(SKILL_DIR))

def generate_word_templates():
    print("=" * 50)
    print("生成 Word 模板文件")
    print("=" * 50)
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("❌ 请先安装: pip install python-docx")
        return False
    
    template_dir = SKILL_DIR / "assets" / "templates" / "word"
    template_dir.mkdir(parents=True, exist_ok=True)
    
    templates = [
        ("letter-business.docx", "商务信函"),
        ("meeting-minutes.docx", "会议纪要"),
        ("project-proposal.docx", "项目提案"),
        ("work-report.docx", "工作报告"),
        ("contract-simple.docx", "合同模板"),
        ("resume-professional.docx", "专业简历"),
        ("press-release.docx", "新闻稿"),
        ("invitation-formal.docx", "正式邀请函"),
    ]
    
    for filename, title in templates:
        print(f"  📄 {filename}", end=" ")
        doc = Document()
        
        # 标题居中
        p = doc.add_heading(title, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加占位说明
        doc.add_paragraph()
        doc.add_paragraph(f"此文件为 {title} 模板").italic = True
        doc.add_paragraph("使用模板变量替换 {{variable_name}} 格式的内容")
        doc.add_paragraph()
        doc.add_paragraph("━" * 40)
        doc.add_paragraph()
        
        # 保存
        filepath = template_dir / filename
        doc.save(str(filepath))
        print(f"✓")
    
    print(f"✅ Word 模板: {len(templates)} 个")
    return True

def generate_excel_templates():
    print("\n" + "=" * 50)
    print("生成 Excel 模板文件")
    print("=" * 50)
    
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        print("❌ 请先安装: pip install openpyxl")
        return False
    
    template_dir = SKILL_DIR / "assets" / "templates" / "excel"
    template_dir.mkdir(parents=True, exist_ok=True)
    
    templates = [
        ("financial-statement.xlsx", "财务报表", ["科目", "期初余额", "本期发生", "期末余额"]),
        ("budget-template.xlsx", "预算表", ["项目", "预算金额", "实际支出", "剩余", "进度"]),
        ("project-timeline.xlsx", "项目进度", ["任务", "负责人", "开始", "结束", "状态"]),
        ("inventory-management.xlsx", "库存管理", ["编号", "名称", "规格", "库存", "安全库存"]),
        ("sales-report.xlsx", "销售报表", ["日期", "产品", "数量", "单价", "金额"]),
        ("attendance-tracking.xlsx", "员工考勤", ["姓名", "日期", "上班", "下班", "状态"]),
        ("crm-simple.xlsx", "客户管理", ["客户", "联系人", "电话", "邮箱", "级别"]),
        ("pivot-demo.xlsx", "数据透视", ["类别", "项目", "数值1", "数值2", "数值3"]),
    ]
    
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    
    for filename, title, headers in templates:
        print(f"  📊 {filename}", end=" ")
        wb = Workbook()
        ws = wb.active
        ws.title = title[:10]  # Excel sheet名最多31字符
        
        # 添加标题行
        ws.append(headers)
        
        # 格式化标题行
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # 调整列宽
        for col_num in range(1, len(headers) + 1):
            ws.column_dimensions[chr(64 + col_num) if col_num <= 26 else f"A{chr(64 + col_num - 26)}"].width = 15
        
        # 添加示例数据行
        ws.append(["示例数据"] * len(headers))
        
        filepath = template_dir / filename
        wb.save(str(filepath))
        print(f"✓")
    
    print(f"✅ Excel 模板: {len(templates)} 个")
    return True

def main():
    print("\n" + "=" * 50)
    print("Office Pro Skill - 模板文件生成器")
    print("=" * 50 + "\n")
    
    success = True
    success = generate_word_templates() and success
    success = generate_excel_templates() and success
    
    print("\n" + "=" * 50)
    if success:
        print("🎉 所有模板文件生成完成!")
        print(f"📁 Word 模板: {SKILL_DIR}/assets/templates/word/")
        print(f"📁 Excel 模板: {SKILL_DIR}/assets/templates/excel/")
    else:
        print("❌ 生成失败，请检查依赖安装")
        return 1
    print("=" * 50 + "\n")
    return 0

if __name__ == "__main__":
    sys.exit(main())
