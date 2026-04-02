#!/usr/bin/env python3
"""
Office Pro - Beautiful Template Generator

Generates professional, aesthetically pleasing Word and Excel templates
with modern design, consistent branding, and proper formatting.
"""

from __future__ import annotations

import os
import sys
from datetime import datetime
from pathlib import Path
from typing import List, Tuple, Optional

SKILL_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(SKILL_DIR))

# Modern color palette
COLORS = {
    'primary': '2B579A',      # Professional blue
    'secondary': '5B9BD5',    # Light blue
    'accent': '70AD47',       # Green accent
    'warning': 'FFC000',      # Amber
    'danger': 'C5504B',       # Red
    'text': '333333',         # Dark gray
    'light': 'F2F2F2',        # Light gray
    'white': 'FFFFFF',
    'border': 'D9D9D9',
}


def generate_word_templates() -> bool:
    """Generate beautiful Word document templates"""
    print("=" * 60)
    print("Generating Beautiful Word Templates")
    print("=" * 60)
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        print(f"[ERROR] Missing dependency: {e}")
        print("Install: pip install python-docx")
        return False
    
    template_dir = SKILL_DIR / "assets" / "templates" / "word"
    template_dir.mkdir(parents=True, exist_ok=True)
    
    def set_cell_shading(cell, color: str):
        """Set cell background color"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)
    
    def add_styled_heading(doc, text: str, level: int = 1):
        """Add a styled heading"""
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor.from_string(COLORS['primary'])
            if level == 1:
                run.font.size = Pt(18)
                run.font.bold = True
        return heading
    
    # Template 1: Meeting Minutes (会议纪要)
    def create_meeting_minutes():
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(10.5)
        font.color.rgb = RGBColor.from_string(COLORS['text'])
        
        # Header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run('MEETING MINUTES')
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(COLORS['primary'])
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('会议纪要')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor.from_string(COLORS['secondary'])
        
        doc.add_paragraph()  # Spacing
        
        # Info table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        info_data = [
            ('Meeting Title / 会议主题', '{{meeting_title}}'),
            ('Date / 日期', '{{meeting_date}}'),
            ('Location / 地点', '{{meeting_location}}'),
            ('Chairperson / 主持人', '{{chairperson}}'),
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            set_cell_shading(row.cells[0], COLORS['light'])
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Attendees section
        add_styled_heading(doc, 'Attendees / 与会人员', level=2)
        doc.add_paragraph('{{attendees}}')
        
        # Agenda section
        add_styled_heading(doc, 'Agenda / 议程', level=2)
        doc.add_paragraph('{{agenda}}')
        
        # Discussion section
        add_styled_heading(doc, 'Discussion / 讨论内容', level=2)
        doc.add_paragraph('{{discussion}}')
        
        # Action items
        add_styled_heading(doc, 'Action Items / 行动项', level=2)
        doc.add_paragraph('{{action_items}}')
        
        # Footer line
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run('_' * 50)
        run.font.color.rgb = RGBColor.from_string(COLORS['border'])
        
        footer_text = doc.add_paragraph()
        footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_text.add_run('Secretary / 记录人: {{secretary}}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(COLORS['text'])
        
        return doc
    
    # Template 2: Business Letter (商务信函)
    def create_business_letter():
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        
        # Letterhead
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run('{{company_name}}')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(COLORS['primary'])
        
        addr = doc.add_paragraph()
        addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = addr.add_run('{{company_address}}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(COLORS['text'])
        
        doc.add_paragraph('_' * 60)
        doc.add_paragraph()
        
        # Date and recipient
        doc.add_paragraph('{{date}}')
        doc.add_paragraph()
        doc.add_paragraph('{{recipient_name}}')
        doc.add_paragraph('{{recipient_title}}')
        doc.add_paragraph('{{recipient_company}}')
        doc.add_paragraph('{{recipient_address}}')
        doc.add_paragraph()
        
        # Salutation
        doc.add_paragraph('Dear {{recipient_name}},')
        doc.add_paragraph()
        
        # Body
        doc.add_paragraph('{{letter_body}}')
        doc.add_paragraph()
        
        # Closing
        doc.add_paragraph('Sincerely,')
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph('{{sender_name}}')
        doc.add_paragraph('{{sender_title}}')
        
        return doc
    
    # Template 3: Professional Resume (专业简历)
    def create_resume():
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(10.5)
        
        # Name header
        name = doc.add_paragraph()
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name.add_run('{{name}}')
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(COLORS['primary'])
        
        # Contact info
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = contact.add_run('{{email}} | {{phone}} | {{location}}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(COLORS['text'])
        
        doc.add_paragraph()
        
        # Sections
        sections = [
            ('Professional Summary / 个人简介', '{{summary}}'),
            ('Work Experience / 工作经历', '{{experience}}'),
            ('Education / 教育背景', '{{education}}'),
            ('Skills / 技能', '{{skills}}'),
            ('Projects / 项目经历', '{{projects}}'),
        ]
        
        for title, placeholder in sections:
            # Section header with bottom border
            heading = doc.add_heading(title, level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor.from_string(COLORS['primary'])
                run.font.size = Pt(12)
                run.font.bold = True
            
            doc.add_paragraph(placeholder)
            doc.add_paragraph()
        
        return doc
    
    # Template 4: Project Proposal (项目提案)
    def create_project_proposal():
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)
        
        # Cover page style title
        doc.add_paragraph()
        doc.add_paragraph()
        
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('PROJECT PROPOSAL')
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(COLORS['primary'])
        
        doc.add_paragraph()
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('{{project_name}}')
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor.from_string(COLORS['secondary'])
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Project info
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run('Prepared by: {{prepared_by}}\n').font.size = Pt(11)
        info.add_run('Date: {{date}}\n').font.size = Pt(11)
        info.add_run('Version: {{version}}').font.size = Pt(11)
        
        doc.add_page_break()
        
        # Content sections
        sections = [
            ('Executive Summary / 执行摘要', '{{executive_summary}}'),
            ('Project Background / 项目背景', '{{background}}'),
            ('Objectives / 项目目标', '{{objectives}}'),
            ('Scope / 项目范围', '{{scope}}'),
            ('Timeline / 时间计划', '{{timeline}}'),
            ('Budget / 预算', '{{budget}}'),
            ('Team / 项目团队', '{{team}}'),
            ('Risk Analysis / 风险分析', '{{risks}}'),
        ]
        
        for title, placeholder in sections:
            heading = doc.add_heading(title, level=1)
            for run in heading.runs:
                run.font.color.rgb = RGBColor.from_string(COLORS['primary'])
            doc.add_paragraph(placeholder)
            doc.add_paragraph()
        
        return doc
    
    # Template 5: Work Report (工作报告)
    def create_work_report():
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(10.5)
        
        # Header
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('WORK REPORT')
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(COLORS['primary'])
        
        doc.add_paragraph()
        
        # Report info table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        info = [
            ('Report Period / 报告期间', '{{report_period}}'),
            ('Department / 部门', '{{department}}'),
            ('Reporter / 报告人', '{{reporter}}'),
        ]
        
        for i, (label, value) in enumerate(info):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            set_cell_shading(row.cells[0], COLORS['light'])
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Sections
        sections = [
            ('Work Summary / 工作总结', '{{summary}}'),
            ('Key Achievements / 主要成果', '{{achievements}}'),
            ('Data Analysis / 数据分析', '{{data_analysis}}'),
            ('Problems & Solutions / 问题与解决', '{{problems}}'),
            ('Next Steps / 下步计划', '{{next_steps}}'),
        ]
        
        for title, placeholder in sections:
            heading = doc.add_heading(title, level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor.from_string(COLORS['primary'])
            doc.add_paragraph(placeholder)
            doc.add_paragraph()
        
        return doc
    
    # Template 6: Contract (合同)
    def create_contract():
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.size = Pt(10.5)
        
        # Contract title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('{{contract_title}}')
        run.font.size = Pt(18)
        run.font.bold = True
        
        doc.add_paragraph()
        
        # Contract number and date
        doc.add_paragraph(f'Contract No.: {{contract_number}}')
        doc.add_paragraph(f'Date: {{contract_date}}')
        doc.add_paragraph()
        
        # Parties
        doc.add_paragraph('Party A (甲方): {{party_a_name}}')
        doc.add_paragraph('Address: {{party_a_address}}')
        doc.add_paragraph()
        doc.add_paragraph('Party B (乙方): {{party_b_name}}')
        doc.add_paragraph('Address: {{party_b_address}}')
        doc.add_paragraph()
        
        # Terms
        doc.add_paragraph('Whereas (鉴于):')
        doc.add_paragraph('{{whereas_clauses}}')
        doc.add_paragraph()
        
        # Articles
        for i in range(1, 9):
            doc.add_paragraph(f'Article {i} / 第{i}条')
            doc.add_paragraph(f'{{article_{i}}}')
            doc.add_paragraph()
        
        # Signature
        doc.add_paragraph()
        doc.add_paragraph('Signature / 签署:')
        doc.add_paragraph()
        
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        table.rows[0].cells[0].text = 'Party A (甲方):\n\n\nSignature: _______________'
        table.rows[0].cells[1].text = 'Party B (乙方):\n\n\nSignature: _______________'
        table.rows[1].cells[0].text = 'Date: _______________'
        table.rows[1].cells[1].text = 'Date: _______________'
        
        return doc
    
    # Template 7: Press Release (新闻稿)
    def create_press_release():
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)
        
        # FOR IMMEDIATE RELEASE
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run('FOR IMMEDIATE RELEASE')
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(COLORS['accent'])
        
        doc.add_paragraph()
        
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('{{headline}}')
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(COLORS['primary'])
        
        doc.add_paragraph()
        
        # Subheadline
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run('{{subheadline}}')
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string(COLORS['secondary'])
        
        doc.add_paragraph()
        doc.add_paragraph('─' * 50)
        doc.add_paragraph()
        
        # Dateline
        dateline = doc.add_paragraph()
        run = dateline.add_run('{{location}} — {{release_date}} — ')
        run.font.bold = True
        
        # Body
        doc.add_paragraph('{{body_paragraph_1}}')
        doc.add_paragraph('{{body_paragraph_2}}')
        doc.add_paragraph('{{body_paragraph_3}}')
        doc.add_paragraph()
        
        # Quote
        quote = doc.add_paragraph()
        quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = quote.add_run('"{{quote}}"')
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string(COLORS['secondary'])
        
        attribution = doc.add_paragraph()
        attribution.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = attribution.add_run('— {{quote_attribution}}')
        run.font.size = Pt(9)
        
        doc.add_paragraph()
        doc.add_paragraph('{{body_paragraph_4}}')
        doc.add_paragraph()
        
        # About section
        doc.add_paragraph('─' * 50)
        heading = doc.add_heading('About {{company_name}}', level=3)
        for run in heading.runs:
            run.font.color.rgb = RGBColor.from_string(COLORS['primary'])
        doc.add_paragraph('{{company_description}}')
        doc.add_paragraph()
        
        # Contact
        heading = doc.add_heading('Media Contact', level=3)
        for run in heading.runs:
            run.font.color.rgb = RGBColor.from_string(COLORS['primary'])
        doc.add_paragraph('{{contact_name}}')
        doc.add_paragraph('{{contact_email}}')
        doc.add_paragraph('{{contact_phone}}')
        
        return doc
    
    # Template 8: Formal Invitation (正式邀请函)
    def create_invitation():
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.size = Pt(12)
        
        # Decorative header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run('◆ ◆ ◆')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor.from_string(COLORS['accent'])
        
        doc.add_paragraph()
        
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('INVITATION')
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(COLORS['primary'])
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('邀请函')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor.from_string(COLORS['secondary'])
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Salutation
        doc.add_paragraph('Dear {{recipient_name}},')
        doc.add_paragraph()
        
        # Invitation text
        doc.add_paragraph(
            '{{host_name}} cordially invites you to attend:'
        )
        doc.add_paragraph()
        
        # Event details box
        event = doc.add_paragraph()
        event.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = event.add_run('{{event_name}}')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(COLORS['primary'])
        
        doc.add_paragraph()
        
        # Details
        details = [
            ('Date / 日期', '{{event_date}}'),
            ('Time / 时间', '{{event_time}}'),
            ('Venue / 地点', '{{event_venue}}'),
            ('Address / 地址', '{{event_address}}'),
        ]
        
        for label, value in details:
            p = doc.add_paragraph()
            p.add_run(f'{label}: ').bold = True
            p.add_run(value)
        
        doc.add_paragraph()
        doc.add_paragraph('{{event_description}}')
        doc.add_paragraph()
        
        # RSVP
        doc.add_paragraph('Please RSVP by {{rsvp_date}}:')
        doc.add_paragraph('{{rsvp_contact}}')
        doc.add_paragraph()
        
        # Closing
        doc.add_paragraph('We look forward to your presence.')
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph('Sincerely,')
        doc.add_paragraph()
        doc.add_paragraph('{{host_name}}')
        doc.add_paragraph('{{host_title}}')
        
        # Footer decoration
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run('◆ ◆ ◆')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor.from_string(COLORS['accent'])
        
        return doc
    
    # Generate all templates
    templates = [
        ('meeting-minutes.docx', create_meeting_minutes),
        ('letter-business.docx', create_business_letter),
        ('resume-professional.docx', create_resume),
        ('project-proposal.docx', create_project_proposal),
        ('work-report.docx', create_work_report),
        ('contract-simple.docx', create_contract),
        ('press-release.docx', create_press_release),
        ('invitation-formal.docx', create_invitation),
    ]
    
    for filename, create_func in templates:
        print(f"  [Word] {filename}...", end=" ")
        try:
            doc = create_func()
            filepath = template_dir / filename
            doc.save(str(filepath))
            print("OK")
        except Exception as e:
            print(f"FAILED: {e}")
            return False
    
    print(f"  Generated {len(templates)} Word templates")
    return True


def generate_excel_templates() -> bool:
    """Generate beautiful Excel spreadsheet templates"""
    print("\n" + "=" * 60)
    print("Generating Beautiful Excel Templates")
    print("=" * 60)
    
    try:
        from openpyxl import Workbook
        from openpyxl.styles import (
            Font, PatternFill, Alignment, Border, Side,
            NamedStyle
        )
        from openpyxl.utils import get_column_letter
        from openpyxl.chart import BarChart, PieChart, Reference
    except ImportError as e:
        print(f"[ERROR] Missing dependency: {e}")
        print("Install: pip install openpyxl")
        return False
    
    template_dir = SKILL_DIR / "assets" / "templates" / "excel"
    template_dir.mkdir(parents=True, exist_ok=True)
    
    # Common styles
    def apply_header_style(cell):
        cell.fill = PatternFill(start_color=COLORS['primary'], 
                                end_color=COLORS['primary'], 
                                fill_type='solid')
        cell.font = Font(color=COLORS['white'], bold=True, size=11)
        cell.alignment = Alignment(horizontal='center', vertical='center')
    
    def apply_subheader_style(cell):
        cell.fill = PatternFill(start_color=COLORS['secondary'], 
                                end_color=COLORS['secondary'], 
                                fill_type='solid')
        cell.font = Font(color=COLORS['white'], bold=True, size=10)
    
    def apply_data_style(cell):
        cell.alignment = Alignment(horizontal='left', vertical='center')
        cell.font = Font(size=10)
    
    thin_border = Border(
        left=Side(style='thin', color=COLORS['border']),
        right=Side(style='thin', color=COLORS['border']),
        top=Side(style='thin', color=COLORS['border']),
        bottom=Side(style='thin', color=COLORS['border'])
    )
    
    # Template 1: Sales Report (销售报表)
    def create_sales_report():
        wb = Workbook()
        ws = wb.active
        ws.title = "Sales Report"
        
        # Title
        ws.merge_cells('A1:F1')
        title_cell = ws['A1']
        title_cell.value = "SALES REPORT / 销售报表"
        title_cell.font = Font(size=16, bold=True, color=COLORS['primary'])
        title_cell.alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 30
        
        # Report info
        ws['A2'] = "Period:"
        ws['B2'] = "{{report_period}}"
        ws['D2'] = "Date:"
        ws['E2'] = "{{report_date}}"
        for cell in [ws['A2'], ws['D2']]:
            cell.font = Font(bold=True, size=10)
        
        # Summary section
        ws['A4'] = "SUMMARY / 汇总"
        ws['A4'].font = Font(bold=True, size=12, color=COLORS['primary'])
        
        summary_headers = ['Metric', 'Value', 'vs Last Period']
        for col, header in enumerate(summary_headers, 1):
            cell = ws.cell(row=5, column=col, value=header)
            apply_header_style(cell)
        
        summary_data = [
            ['Total Sales / 总销售额', '{{total_sales}}', '{{sales_change}}'],
            ['Total Units / 总销量', '{{total_units}}', '{{units_change}}'],
            ['Average Price / 均价', '{{avg_price}}', '{{price_change}}'],
            ['Growth Rate / 增长率', '{{growth_rate}}', ''],
        ]
        
        for row_idx, row_data in enumerate(summary_data, 6):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                apply_data_style(cell)
                if col_idx == 2:  # Value column
                    cell.number_format = '#,##0.00'
        
        # Products section
        ws['A11'] = "PRODUCTS / 产品明细"
        ws['A11'].font = Font(bold=True, size=12, color=COLORS['primary'])
        
        product_headers = ['Product', 'Category', 'Units', 'Price', 'Revenue', 'Share']
        for col, header in enumerate(product_headers, 1):
            cell = ws.cell(row=12, column=col, value=header)
            apply_header_style(cell)
        
        # Sample product row
        ws['A13'] = "{{product_name}}"
        ws['B13'] = "{{category}}"
        ws['C13'] = "{{units_sold}}"
        ws['D13'] = "{{unit_price}}"
        ws['E13'] = "{{revenue}}"
        ws['F13'] = "{{market_share}}"
        
        # Column widths
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 12
        ws.column_dimensions['D'].width = 12
        ws.column_dimensions['E'].width = 15
        ws.column_dimensions['F'].width = 12
        
        return wb
    
    # Template 2: Financial Statement (财务报表)
    def create_financial_statement():
        wb = Workbook()
        ws = wb.active
        ws.title = "Financial"
        
        # Title
        ws.merge_cells('A1:E1')
        ws['A1'] = "FINANCIAL STATEMENT / 财务报表"
        ws['A1'].font = Font(size=16, bold=True, color=COLORS['primary'])
        ws['A1'].alignment = Alignment(horizontal='center')
        ws.row_dimensions[1].height = 30
        
        ws['A2'] = "Period: {{period}}"
        ws['A3'] = "Company: {{company_name}}"
        
        # Balance Sheet
        ws['A5'] = "BALANCE SHEET / 资产负债表"
        ws['A5'].font = Font(bold=True, size=12, color=COLORS['primary'])
        
        headers = ['Item / 项目', 'Opening / 期初', 'Period / 本期', 'Closing / 期末', 'Notes']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=6, column=col, value=header)
            apply_header_style(cell)
        
        # Assets
        ws['A7'] = "ASSETS / 资产"
        ws['A7'].font = Font(bold=True, color=COLORS['secondary'])
        
        asset_items = ['Cash', 'Accounts Receivable', 'Inventory', 'Fixed Assets']
        for idx, item in enumerate(asset_items, 8):
            ws.cell(row=idx, column=1, value=item)
            ws.cell(row=idx, column=2, value="{{" + item.lower().replace(' ', '_') + "_opening}}")
            ws.cell(row=idx, column=3, value="{{" + item.lower().replace(' ', '_') + "_period}}")
            ws.cell(row=idx, column=4, value="{{" + item.lower().replace(' ', '_') + "_closing}}")
        
        # Column widths
        for col in ['A', 'B', 'C', 'D', 'E']:
            ws.column_dimensions[col].width = 18
        
        return wb
    
    # Template 3: Budget Template (预算表)
    def create_budget_template():
        wb = Workbook()
        ws = wb.active
        ws.title = "Budget"
        
        # Title
        ws.merge_cells('A1:F1')
        ws['A1'] = "BUDGET PLAN / 预算计划"
        ws['A1'].font = Font(size=16, bold=True, color=COLORS['primary'])
        ws['A1'].alignment = Alignment(horizontal='center')
        ws.row_dimensions[1].height = 30
        
        ws['A2'] = "Department: {{department}}"
        ws['A3'] = "Fiscal Year: {{fiscal_year}}"
        
        headers = ['Category', 'Budget', 'Actual', 'Variance', '%', 'Status']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=5, column=col, value=header)
            apply_header_style(cell)
        
        # Sample row
        ws['A6'] = "{{category}}"
        ws['B6'] = "{{budget_amount}}"
        ws['C6'] = "{{actual_amount}}"
        ws['D6'] = "{{variance}}"
        ws['E6'] = "{{percentage}}"
        ws['F6'] = "{{status}}"
        
        for col in ['A', 'B', 'C', 'D', 'E', 'F']:
            ws.column_dimensions[col].width = 15
        
        return wb
    
    # Template 4: Project Timeline (项目进度)
    def create_project_timeline():
        wb = Workbook()
        ws = wb.active
        ws.title = "Timeline"
        
        ws.merge_cells('A1:G1')
        ws['A1'] = "PROJECT TIMELINE / 项目进度"
        ws['A1'].font = Font(size=16, bold=True, color=COLORS['primary'])
        ws['A1'].alignment = Alignment(horizontal='center')
        
        ws['A2'] = "Project: {{project_name}}"
        
        headers = ['Task', 'Owner', 'Start', 'End', 'Duration', 'Progress', 'Status']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=4, column=col, value=header)
            apply_header_style(cell)
        
        ws['A5'] = "{{task_name}}"
        ws['B5'] = "{{owner}}"
        ws['C5'] = "{{start_date}}"
        ws['D5'] = "{{end_date}}"
        ws['E5'] = "{{duration}}"
        ws['F5'] = "{{progress}}"
        ws['G5'] = "{{status}}"
        
        for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
            ws.column_dimensions[col].width = 14
        
        return wb
    
    # Template 5: Inventory Management (库存管理)
    def create_inventory():
        wb = Workbook()
        ws = wb.active
        ws.title = "Inventory"
        
        ws.merge_cells('A1:H1')
        ws['A1'] = "INVENTORY MANAGEMENT / 库存管理"
        ws['A1'].font = Font(size=16, bold=True, color=COLORS['primary'])
        ws['A1'].alignment = Alignment(horizontal='center')
        
        headers = ['SKU', 'Name', 'Spec', 'Qty', 'Unit', 'Location', 'Min Stock', 'Status']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col, value=header)
            apply_header_style(cell)
        
        ws['A4'] = "{{sku}}"
        ws['B4'] = "{{product_name}}"
        ws['C4'] = "{{specification}}"
        ws['D4'] = "{{quantity}}"
        ws['E4'] = "{{unit}}"
        ws['F4'] = "{{location}}"
        ws['G4'] = "{{min_stock}}"
        ws['H4'] = "{{stock_status}}"
        
        for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']:
            ws.column_dimensions[col].width = 12
        
        return wb
    
    # Template 6: Attendance Tracking (考勤)
    def create_attendance():
        wb = Workbook()
        ws = wb.active
        ws.title = "Attendance"
        
        ws.merge_cells('A1:H1')
        ws['A1'] = "ATTENDANCE TRACKING / 考勤记录"
        ws['A1'].font = Font(size=16, bold=True, color=COLORS['primary'])
        ws['A1'].alignment = Alignment(horizontal='center')
        
        ws['A2'] = "Month: {{month}}"
        
        headers = ['Date', 'Name', 'ID', 'Check In', 'Check Out', 'Hours', 'Overtime', 'Status']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=4, column=col, value=header)
            apply_header_style(cell)
        
        ws['A5'] = "{{date}}"
        ws['B5'] = "{{employee_name}}"
        ws['C5'] = "{{employee_id}}"
        ws['D5'] = "{{check_in}}"
        ws['E5'] = "{{check_out}}"
        ws['F5'] = "{{work_hours}}"
        ws['G5'] = "{{overtime}}"
        ws['H5'] = "{{attendance_status}}"
        
        for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']:
            ws.column_dimensions[col].width = 12
        
        return wb
    
    # Template 7: CRM Simple (客户管理)
    def create_crm():
        wb = Workbook()
        ws = wb.active
        ws.title = "CRM"
        
        ws.merge_cells('A1:H1')
        ws['A1'] = "CUSTOMER MANAGEMENT / 客户管理"
        ws['A1'].font = Font(size=16, bold=True, color=COLORS['primary'])
        ws['A1'].alignment = Alignment(horizontal='center')
        
        headers = ['Company', 'Contact', 'Phone', 'Email', 'Level', 'Source', 'Last Contact', 'Value']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col, value=header)
            apply_header_style(cell)
        
        ws['A4'] = "{{company_name}}"
        ws['B4'] = "{{contact_name}}"
        ws['C4'] = "{{phone}}"
        ws['D4'] = "{{email}}"
        ws['E4'] = "{{customer_level}}"
        ws['F4'] = "{{lead_source}}"
        ws['G4'] = "{{last_contact_date}}"
        ws['H4'] = "{{lifetime_value}}"
        
        for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']:
            ws.column_dimensions[col].width = 14
        
        return wb
    
    # Template 8: Pivot Demo (数据透视)
    def create_pivot_demo():
        wb = Workbook()
        ws = wb.active
        ws.title = "Data"
        
        ws.merge_cells('A1:E1')
        ws['A1'] = "DATA ANALYSIS / 数据分析"
        ws['A1'].font = Font(size=16, bold=True, color=COLORS['primary'])
        ws['A1'].alignment = Alignment(horizontal='center')
        
        headers = ['Category', 'Item', 'Q1', 'Q2', 'Q3', 'Q4', 'Total']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col, value=header)
            apply_header_style(cell)
        
        ws['A4'] = "{{category}}"
        ws['B4'] = "{{item}}"
        ws['C4'] = "{{q1_value}}"
        ws['D4'] = "{{q2_value}}"
        ws['E4'] = "{{q3_value}}"
        ws['F4'] = "{{q4_value}}"
        ws['G4'] = "{{total}}"
        
        for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G']:
            ws.column_dimensions[col].width = 12
        
        return wb
    
    # Generate all templates
    templates = [
        ('sales-report.xlsx', create_sales_report),
        ('financial-statement.xlsx', create_financial_statement),
        ('budget-template.xlsx', create_budget_template),
        ('project-timeline.xlsx', create_project_timeline),
        ('inventory-management.xlsx', create_inventory),
        ('attendance-tracking.xlsx', create_attendance),
        ('crm-simple.xlsx', create_crm),
        ('pivot-demo.xlsx', create_pivot_demo),
    ]
    
    for filename, create_func in templates:
        print(f"  [Excel] {filename}...", end=" ")
        try:
            wb = create_func()
            filepath = template_dir / filename
            wb.save(str(filepath))
            print("OK")
        except Exception as e:
            print(f"FAILED: {e}")
            return False
    
    print(f"  Generated {len(templates)} Excel templates")
    return True


def main():
    """Main entry point"""
    print("\n" + "=" * 60)
    print("Office Pro - Beautiful Template Generator")
    print("=" * 60 + "\n")
    
    success = True
    success = generate_word_templates() and success
    success = generate_excel_templates() and success
    
    print("\n" + "=" * 60)
    if success:
        print("All templates generated successfully!")
        print(f"Word templates: {SKILL_DIR}/assets/templates/word/")
        print(f"Excel templates: {SKILL_DIR}/assets/templates/excel/")
    else:
        print("Generation failed. Check dependencies.")
        return 1
    print("=" * 60 + "\n")
    return 0


if __name__ == "__main__":
    sys.exit(main())
