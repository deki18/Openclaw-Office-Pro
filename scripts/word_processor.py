"""
Office Pro - Enterprise Document Automation Suite
Word Processor Module

基于 python-docx 和 docxtpl 的企业级 Word 文档处理
"""

from __future__ import annotations

import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from docxtpl import DocxTemplate, InlineImage, RichText
    DOCXTPL_AVAILABLE = True
except ImportError:
    DOCXTPL_AVAILABLE = False


class WordProcessor:
    """
    企业级 Word 文档处理器
    
    支持功能：
    - 文档创建与编辑
    - 模板渲染（Jinja2）
    - 样式管理
    - 表格与图片
    - 页眉页脚
    """
    
    def __init__(self, template_dir: Optional[str] = None):
        """
        初始化 Word 处理器
        
        Args:
            template_dir: 模板目录路径，默认使用内置模板
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required. Install with: pip install python-docx"
            )
        
        self.template_dir = template_dir or self._get_default_template_dir()
        self._document: Optional[DocumentType] = None
        self._template: Optional[DocxTemplate] = None
    
    def _get_default_template_dir(self) -> str:
        """获取默认模板目录"""
        # 假设技能安装在标准位置
        skill_root = Path(__file__).parent.parent
        templates_dir = skill_root / "assets" / "templates" / "word"
        return str(templates_dir)
    
    # ==================== 文档创建与加载 ====================
    
    def create_document(self) -> DocumentType:
        """
        创建新的空白文档
        
        Returns:
            Document 对象
        """
        self._document = Document()
        self._template = None
        return self._document
    
    def load_document(self, path: str) -> DocumentType:
        """
        加载现有文档
        
        Args:
            path: 文档路径
            
        Returns:
            Document 对象
        """
        self._document = Document(path)
        self._template = None
        return self._document
    
    def load_template(self, template_name: str) -> DocxTemplate:
        """
        加载模板文件
        
        Args:
            template_name: 模板文件名（如 'meeting-minutes.docx'）
            
        Returns:
            DocxTemplate 对象
        """
        if not DOCXTPL_AVAILABLE:
            raise ImportError(
                "docxtpl is required for template rendering. "
                "Install with: pip install docxtpl"
            )
        
        template_path = Path(self.template_dir) / template_name
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        self._template = DocxTemplate(str(template_path))
        self._document = self._template.docx
        return self._template
    
    # ==================== 模板渲染 ====================
    
    def render_template(self, context: Dict[str, Any]) -> DocumentType:
        """
        使用上下文数据渲染模板
        
        Args:
            context: 模板变量字典
            
        Returns:
            渲染后的 Document 对象
        """
        if not self._template:
            raise RuntimeError("No template loaded. Call load_template() first.")
        
        self._template.render(context)
        self._document = self._template.docx
        return self._document
    
    def render_and_save(self, context: Dict[str, Any], output_path: str) -> str:
        """
        渲染模板并保存到文件
        
        Args:
            context: 模板变量字典
            output_path: 输出文件路径
            
        Returns:
            保存的文件路径
        """
        self.render_template(context)
        self.save(output_path)
        return output_path
    
    # ==================== 文档操作 ====================
    
    def save(self, path: str) -> None:
        """
        保存文档到指定路径
        
        Args:
            path: 保存路径
        """
        if not self._document:
            raise RuntimeError("No document to save. Create or load a document first.")
        
        # 确保目录存在
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        self._document.save(path)
    
    def add_heading(self, text: str, level: int = 1) -> Any:
        """
        添加标题
        
        Args:
            text: 标题文本
            level: 标题级别（1-9）
            
        Returns:
            Paragraph 对象
        """
        if not self._document:
            raise RuntimeError("No document. Create or load a document first.")
        return self._document.add_heading(text, level=level)
    
    def add_paragraph(self, text: str = "", style: Optional[str] = None) -> Any:
        """
        添加段落
        
        Args:
            text: 段落文本
            style: 样式名称
            
        Returns:
            Paragraph 对象
        """
        if not self._document:
            raise RuntimeError("No document. Create or load a document first.")
        return self._document.add_paragraph(text, style=style)
    
    def add_table(self, rows: int, cols: int, style: Optional[str] = None) -> Any:
        """
        添加表格
        
        Args:
            rows: 行数
            cols: 列数
            style: 表格样式名称
            
        Returns:
            Table 对象
        """
        if not self._document:
            raise RuntimeError("No document. Create or load a document first.")
        
        table = self._document.add_table(rows=rows, cols=cols)
        if style:
            table.style = style
        return table
    
    def add_picture(self, image_path: str, width: Optional[float] = None, height: Optional[float] = None) -> Any:
        """
        添加图片
        
        Args:
            image_path: 图片路径
            width: 宽度（英寸）
            height: 高度（英寸）
            
        Returns:
            InlineShape 对象
        """
        if not self._document:
            raise RuntimeError("No document. Create or load a document first.")
        
        kwargs = {}
        if width:
            kwargs['width'] = Inches(width)
        if height:
            kwargs['height'] = Inches(height)
        
        return self._document.add_picture(image_path, **kwargs)
    
    def add_page_break(self) -> None:
        """添加分页符"""
        if not self._document:
            raise RuntimeError("No document. Create or load a document first.")
        self._document.add_page_break()
    
    # ==================== 页眉页脚 ====================
    
    def add_header(self, text: str, align: str = "center") -> None:
        """
        添加页眉
        
        Args:
            text: 页眉文本
            align: 对齐方式（left/center/right）
        """
        if not self._document:
            raise RuntimeError("No document. Create or load a document first.")
        
        section = self._document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = text
        
        # 设置对齐
        if align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def add_footer(self, text: str, align: str = "center") -> None:
        """
        添加页脚
        
        Args:
            text: 页脚文本
            align: 对齐方式（left/center/right）
        """
        if not self._document:
            raise RuntimeError("No document. Create or load a document first.")
        
        section = self._document.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.text = text
        
        # 设置对齐
        if align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def add_page_number(self, location: str = "footer", align: str = "center") -> None:
        """
        添加页码
        
        Args:
            location: 位置（header/footer）
            align: 对齐方式（left/center/right）
        """
        if not self._document:
            raise RuntimeError("No document. Create or load a document first.")
        
        section = self._document.sections[0]
        
        if location == "header":
            target = section.header
        else:
            target = section.footer
        
        # 获取或创建段落
        if target.paragraphs:
            paragraph = target.paragraphs[0]
        else:
            paragraph = target.add_paragraph()
        
        # 清除现有内容
        paragraph.clear()
        
        # 添加页码字段
        run = paragraph.add_run()
        
        # 创建页码字段（简单文本方式）
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        # 设置对齐
        if align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # ==================== 实用工具 ====================
    
    def get_document_info(self) -> Dict[str, Any]:
        """
        获取文档信息
        
        Returns:
            文档信息字典
        """
        if not self._document:
            raise RuntimeError("No document. Create or load a document first.")
        
        core_props = self._document.core_properties
        
        return {
            'title': core_props.title,
            'author': core_props.author,
            'subject': core_props.subject,
            'keywords': core_props.keywords,
            'created': core_props.created,
            'modified': core_props.modified,
            'paragraph_count': len(self._document.paragraphs),
            'table_count': len(self._document.tables),
        }
    
    @property
    def document(self) -> Optional[DocumentType]:
        """获取当前文档对象"""
        return self._document
    
    @property
    def template(self) -> Optional[DocxTemplate]:
        """获取当前模板对象"""
        return self._template
